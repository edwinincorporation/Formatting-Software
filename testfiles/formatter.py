import sys
import json
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Pure Logic-Based Book Formatter ──────────────────────────

def set_font_properly(run, font_name):
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    # Force for all character sets (Hindi/English)
    for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
        rFonts.set(qn(f'w:{attr}'), font_name)

def detect_structure(text, index):
    """Typesetting Heuristics: Identifies elements based on patterns."""
    clean = text.strip()
    words = clean.split()
    word_count = len(words)
    
    if word_count == 0: return "body"
    if word_count > 15: return "body" # Titles/Chapters are never long paragraphs

    # 1. Main Title: Very first few lines, short, no dots
    if index < 4 and word_count < 8 and not clean.endswith('.'):
        return "title"

    # 2. Chapter Markers: Chapter 1, Unit 2, Adhyaay 5, etc.
    chapter_regex = r'^(chapter|unit|part|section|lesson|adhyaay|अध्याय|इकाई|भाग|vishwa|विश्व)\s*([\dIVX]+)?'
    if re.match(chapter_regex, clean.lower()):
        return "chapter_title"

    # 3. Floating Subheadings: "1.1 Introduction", "2. Basics"
    if re.match(r'^\d+(\.\d+)*\s+', clean):
        return "subheading"

    # 4. Short lines (Upper Case) often are headers
    if clean.isupper() and word_count < 6:
        return "chapter_title"

    return "body"

def apply_clean_justify(para):
    """Implements High-End Typography Justification."""
    text = para.text.strip()
    if len(text.split()) < 15 or len(text) < 120 or text.endswith(('?', ':', '!')):
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return
    
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Remove MS Word's 'Expand' behavior on the last line
    pPr = para._p.get_or_add_pPr()
    for jc in pPr.findall(qn('w:jc')): pPr.remove(jc)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)

def format_document(input_file, output_file, opts):
    doc = Document(input_file)
    font_name = opts.get('font_style') or "Garamond"
    black = RGBColor(0,0,0)

    # 1. Setup Page (Professional Book Standards)
    for section in doc.sections:
        section.page_width, section.page_height = Mm(210), Mm(297) # A4
        section.top_margin = section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.3) # Gutter Margin
        section.right_margin = Inches(0.7)

    # 2. Iterate Paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text: continue

        etype = detect_structure(text, i)

        # Common spacing
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(8)

        if etype == "title":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(72)
            para.paragraph_format.space_after = Pt(36)
            for run in para.runs:
                run.bold = True
                set_font_properly(run, font_name)
                run.font.size, run.font.color.rgb = Pt(28), black

        elif etype == "chapter_title":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(48)
            para.paragraph_format.space_after = Pt(24)
            for run in para.runs:
                run.bold = True
                set_font_properly(run, font_name)
                run.font.size, run.font.color.rgb = Pt(20), black

        elif etype == "subheading":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(14)
            for run in para.runs:
                run.bold = True
                set_font_properly(run, font_name)
                run.font.size, run.font.color.rgb = Pt(14), black

        else:
            # Body Text
            apply_clean_justify(para)
            # Indent first line for books (0.3 inches)
            if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                para.paragraph_format.first_line_indent = Inches(0.3)
            
            for run in para.runs:
                set_font_properly(run, font_name)
                run.font.size, run.font.color.rgb = Pt(12), black

    # 3. Headers/Footers
    for section in doc.sections:
        if opts.get('header'):
            p = section.header.paragraphs[0]
            p.text = opts['header']
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs: set_font_properly(run, font_name); run.font.size = Pt(9)
        
        if opts.get('page_numbers'):
            p = section.footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("Page ")
            set_font_properly(run, font_name); run.font.size = Pt(9)
            # Add field codes for page numbers
            fBegin = OxmlElement('w:fldChar'); fBegin.set(qn('w:fldCharType'), 'begin'); run._r.append(fBegin)
            instr = OxmlElement('w:instrText'); instr.text = ' PAGE '; run._r.append(instr)
            fEnd = OxmlElement('w:fldChar'); fEnd.set(qn('w:fldCharType'), 'end'); run._r.append(fEnd)

    doc.save(output_file)

if __name__ == "__main__":
    in_p = sys.argv[1]
    out_p = sys.argv[2]
    type_d = sys.argv[3]
    opts_f = sys.argv[4]
    
    options = {}
    if os.path.exists(opts_f):
        with open(opts_f, 'r', encoding='utf-8') as f:
            options = json.load(f)
            
    format_document(in_p, out_p, options)
    print(f"Success: {out_p}")








