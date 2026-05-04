import sys
import json
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PAGE_SIZE_MAP = {
    'A4':     (Mm(210), Mm(297)),
    'A5':     (Mm(148), Mm(210)),
    'A3':     (Mm(297), Mm(420)),
    'Letter': (Mm(215.9), Mm(279.4)),
    'Legal':  (Mm(215.9), Mm(355.6)),
}

# ═══════════════════════════
# PRE-CLEANING
# ═══════════════════════════

def clean_runs_in_para(para):
    for run in para.runs:
        cleaned = run.text.replace('\t', '').replace('\n', ' ')
        cleaned = re.sub(r' {2,}', ' ', cleaned)
        run.text = cleaned


def remove_proof_errors(para):
    """Remove <w:proofErr> elements that cause word splits during run iteration."""
    p = para._p
    for proof_err in p.findall(qn('w:proofErr')):
        p.remove(proof_err)


def merge_runs_in_para(para):
    """Merge adjacent runs with identical formatting to prevent mid-word splits.
    Runs split by different rsidR/rsidRPr but same visual formatting get merged."""
    runs = para.runs
    if len(runs) <= 1:
        return

    i = 0
    while i < len(para.runs) - 1:
        r1 = para.runs[i]
        r2 = para.runs[i + 1]

        # Get formatting signatures to compare
        def fmt_sig(run):
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                return ('', None, None, None, None, None)
            bold = rPr.find(qn('w:b'))
            italic = rPr.find(qn('w:i'))
            sz = rPr.find(qn('w:sz'))
            color = rPr.find(qn('w:color'))
            rFonts = rPr.find(qn('w:rFonts'))
            b_val = None if bold is None else bold.get(qn('w:val'), 'true')
            i_val = None if italic is None else italic.get(qn('w:val'), 'true')
            sz_val = None if sz is None else sz.get(qn('w:val'))
            color_val = None if color is None else color.get(qn('w:val'))
            font_val = None if rFonts is None else rFonts.get(qn('w:ascii'))
            return (b_val, i_val, sz_val, color_val, font_val)

        if fmt_sig(r1) == fmt_sig(r2):
            # Merge r2 text into r1, remove r2
            r1.text = (r1.text or '') + (r2.text or '')
            r2._r.getparent().remove(r2._r)
            # Don't increment i — check next pair from same position
        else:
            i += 1

def is_all_bold(para):
    runs = [r for r in para.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)

def is_bullet_para(para):
    """True if paragraph has a numbering/bullet list marker in XML."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return False
    return pPr.find(qn('w:numPr')) is not None

def apply_bold_before_colon(para, font_name, krutidev_mode):
    """If para has 'Label: rest' pattern or ends with ':', make text before ':' bold."""
    text = para.text
    # Find colon.
    colon_idx = text.find(':')
    if colon_idx <= 0 or colon_idx > 80:
        return  # no colon or label too long — skip

    label = text[:colon_idx + 1]
    rest  = text[colon_idx + 1:]

    # Clear all existing runs, rebuild with bold+normal split
    # Collect formatting from first run (size, color) before clearing
    first_run = para.runs[0] if para.runs else None
    size_pt = None
    if first_run:
        size_pt = first_run.font.size

    for run in list(para.runs):
        run._r.getparent().remove(run._r)

    # Bold run: label
    r_bold = para.add_run(label)
    r_bold.bold = True
    if not krutidev_mode:
        set_font_properly(r_bold, font_name)
        if size_pt:
            r_bold.font.size = size_pt
        r_bold.font.color.rgb = RGBColor(0, 0, 0)

    # Normal run: rest
    if rest:
        r_rest = para.add_run(rest)
        r_rest.bold = False
        if not krutidev_mode:
            set_font_properly(r_rest, font_name)
            if size_pt:
                r_rest.font.size = size_pt
            r_rest.font.color.rgb = RGBColor(0, 0, 0)

# def merge_split_paragraphs(doc):
#     paras = doc.paragraphs
#     merge_indices = []
#     i = 0
#     while i < len(paras) - 1:
#         p1 = paras[i]
#         p2 = paras[i + 1]
#         t1 = p1.text.strip()
#         t2 = p2.text.strip()
#         if t1 and t2:
#             b1 = is_all_bold(p1)
#             b2 = is_all_bold(p2)
#             if b1 == b2 and t1[-1].isalpha() and t2[0].islower():
#                 merge_indices.append(i)
#                 i += 2
#                 continue
#         i += 1
#     for idx in reversed(merge_indices):
#         p1 = doc.paragraphs[idx]
#         p2 = doc.paragraphs[idx + 1]
#         for run in p2.runs:
#             p1._p.append(run._r)
#         p2._element.getparent().remove(p2._element)


def merge_split_paragraphs(doc):
    pass  # Disabled — causes duplicate rendering with mixed fonts




def clear_pPr_sz(para):
    """Remove any font size override from paragraph-level rPr (pPr > rPr > sz).
    This prevents paragraph-default sz from overriding run-level sz."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        return
    for tag in [qn('w:sz'), qn('w:szCs')]:
        el = rPr.find(tag)
        if el is not None:
            rPr.remove(el)

def set_pPr_sz(para, half_pts):
    """Set font size at paragraph-level rPr so it applies as default for all runs."""
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    for tag_name in ['w:sz', 'w:szCs']:
        el = rPr.find(qn(tag_name))
        if el is None:
            el = OxmlElement(tag_name)
            rPr.append(el)
        el.set(qn('w:val'), str(half_pts))

def preprocess_document(doc):
    for para in doc.paragraphs:
        remove_proof_errors(para)   # Remove proofErr elements FIRST (they split runs)
        clean_runs_in_para(para)
        merge_runs_in_para(para)    # Merge adjacent same-format runs
    merge_split_paragraphs(doc)

# ═══════════════════════════
# HELPERS
# ═══════════════════════════

KRUTIDEV_FONTS = {'Kruti Dev 010', 'Kruti Dev 011', 'Krutidev010', 'Krutidev011',
                  'KrutiDev010', 'KrutiDev011', 'Kruti Dev010', 'Kruti Dev011'}

# Mapping UI values to exact Windows font names
FONT_NAME_MAP = {
    'Krutidev010': 'Kruti Dev 010',
    'Krutidev011': 'Kruti Dev 011',
    'Mangal': 'Mangal',
    'Kokila': 'Kokila',
    'Utsaah': 'Utsaah',
    'Aparajita': 'Aparajita',
    'Nirmala UI': 'Nirmala UI'
}

# ═══════════════════════════
# HINDI CONVERSION (UNICODE TO KRUTIDEV)
# ═══════════════════════════

def unicode_to_krutidev(text):
    if not text: return ""
    
    # Check if text contains any Unicode Hindi characters. If not, return as is.
    if not re.search(r'[\u0900-\u097F]', text):
        return text

    # Standard mapping for Unicode to Kruti Dev 010
    # This is a simplified version of standard transliteration logic
    array_unicode = [
        "क़", "ख़", "ग़", "ज़", "ड़", "ढ़", "फ़", "य़", "ऱ", "ऩ", 
        "ा", "ि", "ी", "ु", "ू", "ृ", "े", "ै", "ो", "ौ", "ं", "ः", "ँ",
        "०", "१", "२", "३", "४", "५", "६", "७", "८", "९",
        "अ", "आ", "इ", "ई", "उ", "ऊ", "ऋ", "ए", "ऐ", "ओ", "औ",
        "क", "ख", "ग", "घ", "ङ",
        "च", "छ", "ज", "झ", "ञ",
        "ट", "ठ", "ड", "ढ", "ण",
        "त", "थ", "द", "ध", "न",
        "प", "फ", "ब", "भ", "म",
        "य", "र", "ल", "व", "श", "ष", "स", "ह",
        "क्ष", "त्र", "ज्ञ", "श्र", "ज्ञ",
        "।", "॥", "्", "‍"
    ]
    
    array_krutidev = [
        "क़", "ख़", "ग़", "ज़", "ड़", "ढ़", "फ़", "य़", "ऱ", "ऩ",
        "k", "f", "h", "q", "w", "=", "s", "S", "ks", "kS", "a", "A", "¡",
        ")", "!", "@", "#", "$", "%", "^", "&", "*", "(",
        "v", "vk", " b", "bZ", "m", "Å", "½", "vks", "vS", "vks", "vS",
        "d", "[k", "x", "?", "³",
        "p", "N", "t", "÷", "¥",
        "V", "B", "M", "<", ".k",
        "r", "Fk", "n", "èk", "u",
        "i", "Q", "c", "Hk", "e",
        "य", "j", "y", "o", "श", "ष", "l", "ह",
        "़", "त्र", "ज्ञ", "श्र", "ज्ञ",
        "P", "Q", "्", ""
    ]

    # This is a complex process. For a robust implementation, we need 
    # specific rules for matras that come before characters (like 'i').
    
    # 1. Simple replacements
    modified_text = text
    
    # Handle nukta characters first
    modified_text = modified_text.replace("क़", "d+").replace("ख़", "[k+").replace("ग़", "x+").replace("ज़", "t+").replace("ड़", "M+").replace("ढ़", "<+").replace("फ़", "i+").replace("य़", "य़")
    
    # Handle composite characters
    modified_text = modified_text.replace("क्ष", "d" + "़").replace("त्र", "=").replace("ज्ञ", "K").replace("श्र", "श" + "्" + "j")
    
    # Character by character map (basic)
    # Note: Full robust conversion requires rule-based reordering for 'i' matra, etc.
    # We will implement the most common ones.
    
    mapping = {
        '।': 'A', '॥': 'AA', '्': '्', 
        'ा': 'k', 'ि': 'f', 'ी': 'h', 'ु': 'q', 'ू': 'w', 'ृ': '=', 'े': 's', 'ै': 'S', 'ो': 'ks', 'ौ': 'kS', 'ं': 'a', 'ः': '%', 'ँ': '¡',
        'अ': 'v', 'आ': 'vk', 'इ': 'b', 'ई': 'bZ', 'उ': 'm', 'ऊ': 'Å', 'ए': 's', 'ऐ': 'S', 'ओ': 'vks', 'औ': 'vS',
        'क': 'd', 'ख': '[k', 'ग': 'x', 'घ': '?', 'ङ': '³',
        'च': 'p', 'छ': 'N', 'ज': 't', 'झ': '÷', 'ञ': '¥',
        'ट': 'V', 'ठ': 'B', 'ड': 'M', 'ढ': '<', 'ण': '.k',
        'त': 'r', 'थ': 'Fk', 'द': 'n', 'ध': 'èk', 'न': 'u',
        'प': 'i', 'फ': 'Q', 'ब': 'c', 'भ': 'Hk', 'म': 'e',
        'य': ';', 'र': 'j', 'ल': 'y', 'व': 'o', 'श': 'श', 'ष': 'ष', 'स': 'l', 'ह': 'g',
        '०': ')', '१': '!', '२': '@', '३': '#', '४': '$', '५': '%', '६': '^', '७': '&', '८': '*', '९': '('
    }

    # Rule-based conversion for Kruti Dev (Complex)
    # 1. Handle 'i' matra (Chhoti ee) which comes BEFORE the consonant in Kruti Dev
    # Regex to find (Consonant)(Matra 'i') and swap them
    modified_text = re.sub(r'([\u0915-\u0939])\u093f', r'f\1', modified_text)
    # Handle half consonants + 'i' matra
    modified_text = re.sub(r'([\u0915-\u0939])\u094d([\u0915-\u0939])\u093f', r'f\1\u094d\2', modified_text)

    # 2. Map characters
    res = ""
    for char in modified_text:
        res += mapping.get(char, char)
    
    # 3. Handle Halant (Half characters)
    # In Kruti Dev, half characters are often different glyphs.
    # d + ् -> d (different key for half k)
    # This part is highly font-specific. For Krutidev010:
    res = res.replace("d्", "D").replace("[k्", "K").replace("x्", "X").replace("?्", "ഘ").replace("p्", "P").replace("t्", "T").replace("r्", "R").replace("Fk्", "f").replace("n्", "N").replace("èk्", "E").replace("u्", "U").replace("i्", "I").replace("Q्", "q").replace("c्", "C").replace("Hk्", "h").replace("e्", "E").replace(";्", "Y").replace("y्", "L").replace("o्", "O").replace("l्", "L")
    
    # Basic cleanup
    res = res.replace("्", "") # Remove remaining halants
    
    return res

def is_krutidev(font_name):
    return font_name and any(k.lower() in font_name.lower() for k in ['kruti', 'krutidev'])

def set_font_properly(run, font_name, size_pt=None):
    # Map to formal name if exists
    formal_name = FONT_NAME_MAP.get(font_name, font_name)
    run.font.name = formal_name
    
    r = run._element
    rPr = r.get_or_add_rPr()
    
    # 1. Set Fonts
    rFonts = rPr.get_or_add_rFonts()
    # Kruti Dev is a legacy Western font. Setting 'hint' to 'default' is crucial.
    if is_krutidev(formal_name):
        rFonts.set(qn('w:hint'), 'default')
    else:
        rFonts.set(qn('w:hint'), 'complex')

    for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
        rFonts.set(qn(f'w:{attr}'), formal_name)
    
    # 2. Set Language (Crucial Fix)
    # If Kruti Dev, we MUST tell Word it is English/Latin so it doesn't fall back to Arial Unicode
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    
    if is_krutidev(formal_name):
        # Force Latin/Western IDs
        lang.set(qn('w:val'), 'en-US')
        lang.set(qn('w:ascii'), 'en-US')
        lang.set(qn('w:hAnsi'), 'en-US')
        lang.set(qn('w:bidi'), 'hi-IN') # Bi-directional as secondary
    else:
        lang.set(qn('w:val'), 'hi-IN')
        lang.set(qn('w:cs'), 'hi-IN')

    # 3. Set Size
    if size_pt:
        run.font.size = Pt(size_pt)
        # Also set size in rPr for complex scripts/CS
        sz_cs = rPr.find(qn('w:szCs'))
        if sz_cs is None:
            sz_cs = OxmlElement('w:szCs')
            rPr.append(sz_cs)
        sz_cs.set(qn('w:val'), str(int(size_pt * 2)))

def set_para_font(para, font_name):
    """Set font at paragraph-level rPr."""
    formal_name = FONT_NAME_MAP.get(font_name, font_name)
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    
    if is_krutidev(formal_name):
        rFonts.set(qn('w:hint'), 'default')
    
    attrs = ['ascii', 'hAnsi', 'eastAsia', 'cs']
    for attr in attrs:
        rFonts.set(qn(f'w:{attr}'), formal_name)
    
    # Set Language at para level too
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    if is_krutidev(formal_name):
        lang.set(qn('w:val'), 'en-US')
    else:
        lang.set(qn('w:val'), 'hi-IN')

def add_run_with_font(para, text, font_name, size_pt, bold=False, color=None):
    run = para.add_run(text)
    run.bold = bold
    set_font_properly(run, font_name)
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run

def add_fld_char(run, fld_type):
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), fld_type)
    run._r.append(fc)

def add_instr_text(run, instr):
    it = OxmlElement('w:instrText')
    it.text = instr
    run._r.append(it)

# ═══════════════════════════
# TITLE PAGE
# ═══════════════════════════

def insert_title_page(doc, opts, font_name):
    black = RGBColor(0, 0, 0)
    gray  = RGBColor(100, 100, 100)
    title       = opts.get('title', '').strip()
    author      = opts.get('author', '').strip()
    volume      = opts.get('volume', '').strip()
    isbn        = opts.get('isbn', '').strip()
    website     = opts.get('website_url', '').strip()
    footer_text = opts.get('footer', '').strip()

    if not title and not author:
        return

    insert_paras = []

    def make_para(text, align, size, bold=False, space_before=0, space_after=12, color=None):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if text:
            add_run_with_font(p, text, font_name, size, bold=bold, color=color or black)
        return p

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(72)
    spacer.paragraph_format.space_after  = Pt(0)
    insert_paras.append(spacer)

    if title:
        insert_paras.append(make_para(title, WD_ALIGN_PARAGRAPH.CENTER, 28, bold=True, space_after=20))
    if volume:
        insert_paras.append(make_para(volume, WD_ALIGN_PARAGRAPH.CENTER, 14, space_after=14, color=gray))
    if author:
        insert_paras.append(make_para(author, WD_ALIGN_PARAGRAPH.CENTER, 16, bold=True, space_before=10, space_after=14))

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_before = Pt(20)
    sep.paragraph_format.space_after  = Pt(20)
    r = sep.add_run('\u2015 \u2015 \u2015')
    set_font_properly(r, font_name)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(180, 180, 180)
    insert_paras.append(sep)

    if footer_text:
        insert_paras.append(make_para(footer_text, WD_ALIGN_PARAGRAPH.CENTER, 11, space_after=8, color=gray))
    if website:
        insert_paras.append(make_para(website, WD_ALIGN_PARAGRAPH.CENTER, 10, space_after=8, color=gray))
    if isbn:
        insert_paras.append(make_para(f'ISBN: {isbn}', WD_ALIGN_PARAGRAPH.CENTER, 10, space_after=0, color=gray))

    pb_para = doc.add_paragraph()
    pb_para.paragraph_format.space_before = Pt(0)
    pb_para.paragraph_format.space_after  = Pt(0)
    run = pb_para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    insert_paras.append(pb_para)

    body = doc.element.body
    for p in reversed(insert_paras):
        body.remove(p._element)
        body.insert(0, p._element)

# ═══════════════════════════
# THESIS TITLE PAGE
# ═══════════════════════════

def insert_thesis_title_page(doc, opts, font_name):
    """Academic thesis title page: University → Title → Student → Supervisor → Year"""
    black = RGBColor(0, 0, 0)

    title       = opts.get('title', '').strip()
    author      = opts.get('author', '').strip()
    university  = opts.get('university', '').strip()
    department  = opts.get('department', '').strip()
    supervisor  = opts.get('supervisor', '').strip()
    year        = opts.get('year', '').strip()

    if not title and not author:
        return

    insert_paras = []

    def make_para(text, align, size, bold=False, italic=False,
                  space_before=0, space_after=10, color=None):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if text:
            r = p.add_run(text)
            r.bold   = bold
            r.italic = italic
            set_font_properly(r, font_name)
            r.font.size = Pt(size)
            r.font.color.rgb = black  # Force black
        return p

    def add_horizontal_rule(thick=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12' if thick else '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')  # Force black
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # Top spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(36)
    spacer.paragraph_format.space_after  = Pt(0)
    insert_paras.append(spacer)

    # University name (large, bold, black)
    if university:
        insert_paras.append(make_para(university, WD_ALIGN_PARAGRAPH.CENTER,
                                      16, bold=True, space_after=4))
    # Department
    if department:
        insert_paras.append(make_para(department, WD_ALIGN_PARAGRAPH.CENTER,
                                      12, space_after=20))

    # Thick rule
    insert_paras.append(add_horizontal_rule(thick=True))

    # Mid spacer
    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_before = Pt(24)
    sp2.paragraph_format.space_after  = Pt(0)
    insert_paras.append(sp2)

    # "A Thesis Submitted for..." label
    insert_paras.append(make_para('A Thesis Submitted in Partial Fulfillment of the',
                                   WD_ALIGN_PARAGRAPH.CENTER, 11, italic=True, space_after=2))
    insert_paras.append(make_para('Requirements for the Degree',
                                   WD_ALIGN_PARAGRAPH.CENTER, 11, italic=True, space_after=24))

    # Title (biggest element)
    if title:
        insert_paras.append(make_para(title, WD_ALIGN_PARAGRAPH.CENTER,
                                      22, bold=True,
                                      space_before=8, space_after=28))

    # Thin rule
    insert_paras.append(add_horizontal_rule(thick=False))

    # Submitted by label
    insert_paras.append(make_para('Submitted by', WD_ALIGN_PARAGRAPH.CENTER,
                                   10, italic=True,
                                   space_before=20, space_after=4))
    # Student name
    if author:
        insert_paras.append(make_para(author, WD_ALIGN_PARAGRAPH.CENTER,
                                      15, bold=True, space_after=4))

    # Supervisor
    if supervisor:
        insert_paras.append(make_para(f'Under the Supervision of', WD_ALIGN_PARAGRAPH.CENTER,
                                       10, italic=True,
                                       space_before=16, space_after=4))
        insert_paras.append(make_para(supervisor, WD_ALIGN_PARAGRAPH.CENTER,
                                       13, bold=True, space_after=4))

    # Year
    if year:
        insert_paras.append(make_para(year, WD_ALIGN_PARAGRAPH.CENTER,
                                       12, space_before=20, space_after=0))

    # Page break
    pb_para = doc.add_paragraph()
    pb_para.paragraph_format.space_before = Pt(0)
    pb_para.paragraph_format.space_after  = Pt(0)
    run = pb_para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    insert_paras.append(pb_para)

    body = doc.element.body
    for p in reversed(insert_paras):
        body.remove(p._element)
        body.insert(0, p._element)


# ═══════════════════════════
# LETTER TITLE PAGE
# ═══════════════════════════

def insert_letter_header(doc, opts, font_name):
    """Letter / Notice header: Org name → Ref+Date → Subject line"""
    black  = RGBColor(0, 0, 0)
    gray   = RGBColor(80, 80, 80)
    dark   = RGBColor(20, 20, 80)

    org_name = opts.get('org_name', '').strip()
    ref_no   = opts.get('ref_no', '').strip()
    date     = opts.get('date', '').strip()
    subject  = opts.get('subject', '').strip()

    if not org_name and not subject:
        return

    insert_paras = []

    def make_para(text, align, size, bold=False, italic=False,
                  space_before=0, space_after=8, color=None):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if text:
            r = p.add_run(text)
            r.bold   = bold
            r.italic = italic
            set_font_properly(r, font_name)
            r.font.size = Pt(size)
            r.font.color.rgb = color or black
        return p

    # Org name — centered, large
    if org_name:
        insert_paras.append(make_para(org_name, WD_ALIGN_PARAGRAPH.CENTER,
                                       16, bold=True, color=dark,
                                       space_before=0, space_after=4))

    # Thin rule under org name
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(4)
    hr.paragraph_format.space_after  = Pt(10)
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2222AA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    insert_paras.append(hr)

    # Ref + Date on same line (left + right)
    if ref_no or date:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        if ref_no:
            r1 = p.add_run(f'Ref.: {ref_no}')
            r1.bold = True
            set_font_properly(r1, font_name)
            r1.font.size = Pt(11)
            r1.font.color.rgb = black

        if ref_no and date:
            # Tab to push date to right
            tab_r = p.add_run('\t\t\t\t\t\t')
            set_font_properly(tab_r, font_name)
            tab_r.font.size = Pt(11)

        if date:
            r2 = p.add_run(f'Date: {date}')
            r2.bold = True
            set_font_properly(r2, font_name)
            r2.font.size = Pt(11)
            r2.font.color.rgb = black

        insert_paras.append(p)

    # Subject line
    if subject:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(12)
        sp.paragraph_format.space_after  = Pt(12)
        sp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        r_lbl = sp.add_run('Subject: ')
        r_lbl.bold = True
        set_font_properly(r_lbl, font_name)
        r_lbl.font.size = Pt(12)
        r_lbl.font.color.rgb = black

        r_sub = sp.add_run(subject)
        r_sub.bold = True
        r_sub.underline = True
        set_font_properly(r_sub, font_name)
        r_sub.font.size = Pt(12)
        r_sub.font.color.rgb = dark

        insert_paras.append(sp)

    # Rule below subject
    hr2 = doc.add_paragraph()
    hr2.paragraph_format.space_before = Pt(4)
    hr2.paragraph_format.space_after  = Pt(16)
    pPr2 = hr2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    b2 = OxmlElement('w:bottom')
    b2.set(qn('w:val'), 'single')
    b2.set(qn('w:sz'), '4')
    b2.set(qn('w:space'), '1')
    b2.set(qn('w:color'), 'AAAACC')
    pBdr2.append(b2)
    pPr2.append(pBdr2)
    insert_paras.append(hr2)

    body = doc.element.body
    for p in reversed(insert_paras):
        body.remove(p._element)
        body.insert(0, p._element)


# ═══════════════════════════
# THESIS BODY FORMATTING
# ═══════════════════════════

def detect_thesis_structure(para, index, doc):
    """Thesis-aware structure detection."""
    text  = para.text.strip()
    words = text.split()
    wc    = len(words)

    if wc == 0:
        return 'empty'
    if is_bullet_para(para):
        return 'bullet'

    is_bold = is_all_bold(para)

    # Chapter detection (matches 'Chapter 8' or 'Chapter 8: Title')
    if text.lower().startswith('chapter') and wc <= 15:
        return 'chapter_heading'
    
    # Check if this paragraph is a chapter title (follows a chapter heading)
    if index > 0:
        prev_para = doc.paragraphs[index-1]
        if prev_para.text.strip().lower().startswith('chapter') and len(prev_para.text.split()) <= 4 and wc <= 10:
            return 'chapter_heading'

    # Abstract / Keywords / References as special headings
    special_sections = ['abstract', 'introduction', 'conclusion', 'references',
                        'bibliography', 'acknowledgement', 'appendix', 'keywords',
                        'methodology', 'literature review', 'discussion', 'results']
    if text.lower().strip('.') in special_sections and wc <= 4:
        return 'section_heading'

    # Numbered section: 1. or 1.1 or 1.1.1
    if re.match(r'^\d+(\.\d+)*[\.\s]', text) and is_bold and wc <= 15:
        return 'subheading'

    # ALL CAPS short = section heading
    if text.isupper() and 1 < wc < 8:
        return 'section_heading'

    # Bold short = subheading
    if is_bold and wc <= 15:
        return 'subheading'

    if wc > 15:
        return 'body'

    return 'body'


def format_thesis_body(doc, opts, font_name):
    """Apply thesis-specific paragraph formatting to body."""
    black  = RGBColor(0, 0, 0)
    krutidev_mode = is_krutidev(font_name)
    base_size = float(opts.get('font_size', 12))
    line_spacing = float(opts.get('line_spacing', 1.15))

    # We need to use a while loop because we might insert paragraphs
    i = 0
    prev_etype = None
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if not text:
            i += 1
            continue

        etype = detect_thesis_structure(para, i, doc)
        if etype == 'empty':
            i += 1
            continue

        # Default spacing rule: 4.0 spacing after every paragraph
        space_after = 4.0
        
        # Look ahead to see if next paragraph is body/bullet
        next_etype = None
        if i < len(doc.paragraphs) - 1:
            next_para = doc.paragraphs[i+1]
            if next_para.text.strip():
                next_etype = detect_thesis_structure(next_para, i+1, doc)
        
        # If heading is followed by body or bullet, reduce space_after
        if etype in ['section_heading', 'subheading'] and next_etype in ['body', 'bullet']:
            space_after = 2.0

        # Reduce space before if the previous paragraph was also a heading
        space_before = 14.0
        if etype in ['section_heading', 'subheading'] and prev_etype in ['chapter_heading', 'section_heading', 'subheading']:
            space_before = 2.0 # Reduced spacing between headings

        if etype == 'chapter_heading':
            # Handle 'Chapter 8: Title' by splitting it
            if ':' in text and text.lower().startswith('chapter'):
                parts = text.split(':', 1)
                chapter_label = parts[0].strip()
                chapter_title = parts[1].strip()
                
                para.text = chapter_label.upper()
                apply_para_formatting(para, etype, font_name,
                    font_size_pt=base_size + 2, bold=True, color=black,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=24, space_after_pt=0,
                    line_spacing=line_spacing)
                
                title_para = doc.add_paragraph(chapter_title)
                para._p.addnext(title_para._p)
                
                apply_para_formatting(title_para, etype, font_name,
                    font_size_pt=base_size + 2, bold=True, color=black,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=0, space_after_pt=space_after,
                    line_spacing=line_spacing)
                i += 2
                prev_etype = 'chapter_heading'
                continue
            else:
                apply_para_formatting(para, etype, font_name,
                    font_size_pt=base_size + 2, bold=True, color=black,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=24, space_after_pt=space_after,
                    line_spacing=line_spacing)

        elif etype == 'section_heading':
            apply_para_formatting(para, etype, font_name,
                font_size_pt=base_size + 1, bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=space_before + 4, space_after_pt=space_after,
                line_spacing=line_spacing)

        elif etype == 'subheading':
            apply_para_formatting(para, etype, font_name,
                font_size_pt=base_size, bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=space_before, space_after_pt=space_after,
                line_spacing=line_spacing)
            if not krutidev_mode and ':' in para.text:
                apply_bold_before_colon(para, font_name, krutidev_mode)

        elif etype == 'bullet':
            is_bold = is_all_bold(para)
            apply_para_formatting(para, etype, font_name,
                font_size_pt=base_size, bold=is_bold, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=0, space_after_pt=space_after,
                line_spacing=line_spacing)

        else:  # body
            apply_clean_justify(para)
            final_align = para.alignment if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY else WD_ALIGN_PARAGRAPH.LEFT
            
            apply_para_formatting(para, etype, font_name,
                font_size_pt=base_size, bold=False, color=black,
                align=final_align,
                space_before_pt=0, space_after_pt=space_after,
                first_indent=None,
                line_spacing=line_spacing)
            
            if ':' in para.text:
                apply_bold_before_colon(para, font_name, krutidev_mode)
        
        prev_etype = etype
        i += 1


# ═══════════════════════════
# LETTER BODY FORMATTING
# ═══════════════════════════

def detect_letter_structure(para, index):
    """Letter-aware structure detection."""
    text  = para.text.strip()
    words = text.split()
    wc    = len(words)

    if wc == 0:
        return 'empty'
    if is_bullet_para(para):
        return 'bullet'

    is_bold = is_all_bold(para)

    # Salutation: Dear ...,
    if re.match(r'^(dear|to|respected|sub|subject)', text.lower()):
        return 'salutation'

    # Closing: Yours sincerely / Best regards / Thanking you
    closing_words = ['yours', 'sincerely', 'faithfully', 'regards', 'thanking',
                     'with regards', 'best regards', 'warm regards']
    if any(text.lower().startswith(w) for w in closing_words) and wc <= 5:
        return 'closing'

    # Signature block (bold, short)
    if is_bold and wc <= 8 and index > 5:
        return 'signature'

    # Short bold lines = label/header within letter
    if is_bold and wc <= 12:
        return 'label'

    return 'body'


def has_existing_letter_header(doc):
    """Return True if document already contains a Ref./Date line (pre-existing letterhead)."""
    for para in doc.paragraphs[:10]:
        t = para.text.strip()
        if re.match(r'^ref\.?\s*:', t, re.IGNORECASE):
            return True
    return False


def is_ref_date_line(para):
    """True if paragraph is the Ref + Date line that must be preserved as-is."""
    return bool(re.match(r'^ref\.?\s*:', para.text.strip(), re.IGNORECASE))


def preserve_para_indent(para):
    """Return copy of existing w:ind element so we can restore it after formatting."""
    import copy
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        return None
    return copy.deepcopy(ind)


def restore_para_indent(para, saved_ind):
    """Restore a previously saved w:ind element to the paragraph."""
    if saved_ind is None:
        return
    pPr = para._p.get_or_add_pPr()
    existing = pPr.find(qn('w:ind'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(saved_ind)


def format_letter_body(doc, opts, font_name):
    """Apply letter-specific paragraph formatting."""
    black  = RGBColor(0, 0, 0)
    dark   = RGBColor(20, 20, 80)
    gray   = RGBColor(80, 80, 80)
    krutidev_mode = is_krutidev(font_name)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Ref+Date line — preserve original indent & spacing, only update font
        if is_ref_date_line(para):
            saved_ind = preserve_para_indent(para)
            if not krutidev_mode:
                set_para_font(para, font_name)
                for run in para.runs:
                    set_font_properly(run, font_name)
                    run.font.size = Pt(11)
            restore_para_indent(para, saved_ind)
            continue

        etype = detect_letter_structure(para, i)
        if etype == 'empty':
            continue

        if etype == 'salutation':
            apply_para_formatting(para, etype, font_name,
                font_size_pt=12, bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=8, space_after_pt=8)

        elif etype == 'closing':
            apply_para_formatting(para, etype, font_name,
                font_size_pt=12, bold=False, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=16, space_after_pt=4)

        elif etype == 'signature':
            apply_para_formatting(para, etype, font_name,
                font_size_pt=12, bold=True, color=dark,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=2, space_after_pt=2)

        elif etype == 'label':
            apply_para_formatting(para, etype, font_name,
                font_size_pt=12, bold=True, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=12, space_after_pt=4)
            if not krutidev_mode and ': ' in para.text:
                apply_bold_before_colon(para, font_name, krutidev_mode)

        elif etype == 'bullet':
            is_bold = is_all_bold(para)
            apply_para_formatting(para, etype, font_name,
                font_size_pt=12, bold=is_bold, color=black,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_before_pt=0, space_after_pt=4)
            if not krutidev_mode and ': ' in para.text and not is_bold:
                apply_bold_before_colon(para, font_name, krutidev_mode)

        else:  # body — single spaced, justified
            if krutidev_mode:
                apply_para_formatting(para, etype, font_name,
                    font_size_pt=12, bold=False, color=black,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before_pt=0, space_after_pt=4)
            else:
                apply_clean_justify(para)
                apply_para_formatting(para, etype, font_name,
                    font_size_pt=12, bold=False, color=black,
                    align=para.alignment,
                    space_before_pt=0, space_after_pt=6)


# ═══════════════════════════
# STRUCTURE DETECTION
# ═══════════════════════════

def get_original_alignment(para):
    """Return the original paragraph alignment from XML (before any override)."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        return None
    val = jc.get(qn('w:val'))
    mapping = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
        'both':   WD_ALIGN_PARAGRAPH.JUSTIFY,
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
    }
    return mapping.get(val)


def detect_structure(para, index, doc=None):
    """Detect paragraph type. Bullet paragraphs are always 'bullet'."""
    text  = para.text.strip()
    words = text.split()
    wc    = len(words)

    if wc == 0:
        return 'empty'

    # Bullet list items — NEVER treat as heading regardless of bold
    if is_bullet_para(para):
        return 'bullet'
    
    # High Priority: The very first non-empty paragraph of a book is often the title/chapter name
    if index == 0 and wc <= 15:
        return 'chapter_title'

    if wc > 20:
        return 'body'

    is_bold = is_all_bold(para)
    orig_align = get_original_alignment(para)

    # Chapter detection (matches 'Chapter 8' or 'Chapter 8: Title' or common Hindi markers)
    chapter_regex = r'^(chapter|unit|part|section|lesson|adhyaay|\u0905\u0927\u094d\u092f\u093e\u092f|\u0907\u0915\u093e\u0908|\u092d\u093e\u0917)\s*([\dIVX]+)?'
    if re.match(chapter_regex, text.lower()) and wc <= 15:
        return 'chapter_title'

    if re.match(r'^\d+(\.\d+)*[\.\s]', text) and is_bold and wc <= 12:
        return 'subheading'

    if re.match(r'^[a-zA-Z]\)', text) and is_bold and wc <= 10:
        return 'subheading'

    if text.isupper() and wc < 8:
        return 'chapter_title'

    if index < 4 and wc < 8 and is_bold and not text.endswith('.'):
        return 'title'

    if is_bold and wc <= 15:
        return 'subheading'

    return 'body'

# ═══════════════════════════
# JUSTIFY
# ═══════════════════════════

def set_font_properly(run, font_name, size_pt=None):
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    # Explicitly set all font attributes. 'cs' is vital for Hindi/Complex Scripts.
    for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
        rFonts.set(qn(f'w:{attr}'), font_name)
    
    if size_pt:
        run.font.size = Pt(size_pt)
        # Also set size in rPr for complex scripts
        sz_cs = rPr.find(qn('w:szCs'))
        if sz_cs is None:
            sz_cs = OxmlElement('w:szCs')
            rPr.append(sz_cs)
        sz_cs.set(qn('w:val'), str(int(size_pt * 2)))

def apply_clean_justify(para):
    """Justify only long lines; short lines stay left-aligned to avoid awkward gaps."""
    text = para.text.strip()
    words = text.split()
    # Stricter thresholds: need 12+ words AND 100+ chars to justify.
    # Also skip if it ends with punctuation that usually marks the end of a short line.
    if len(words) < 12 or len(text) < 100 or text.endswith(('?', ':', '!', ';')):
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return

    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = para._p.get_or_add_pPr()
    for jc in pPr.findall(qn('w:jc')):
        pPr.remove(jc)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)

def apply_para_formatting(para, etype, font_name, font_size_pt, bold, color, align,
                           space_before_pt, space_after_pt, first_indent=None, line_spacing=1.15):
    """Apply all formatting to a paragraph — both run-level and pPr-level."""
    set_para_font(para, font_name)
    clear_pPr_sz(para)
    set_pPr_sz(para, int(font_size_pt * 2))

    # Spacing & Line Spacing
    para.paragraph_format.space_before = Pt(space_before_pt)
    para.paragraph_format.space_after  = Pt(space_after_pt)
    
    # Handle line spacing (1.0, 1.15, 1.5, 2.0)
    try:
        ls = float(line_spacing)
    except:
        ls = 1.15

    if ls == 1.0:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif ls == 2.0:
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        # For 1.15 or 1.5, use MULTIPLE
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.line_spacing = ls

    if first_indent is not None:
        para.paragraph_format.first_line_indent = first_indent
    else:
        para.paragraph_format.first_line_indent = None

    # Alignment
    para.alignment = align

    # Run-level
    for run in para.runs:
        run.bold = bold
        set_font_properly(run, font_name, font_size_pt)
        run.font.color.rgb = color

# ═══════════════════════════
# MAIN
# ═══════════════════════════

def center_all_tables(doc):
    """Center-align every table on the page."""
    from docx.oxml.ns import qn
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        jc = tblPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            tblPr.append(jc)
        jc.set(qn('w:val'), 'center')

def format_document(input_file, output_file, opts, doc_type='book'):
    doc = Document(input_file)
    font_name = opts.get('font_style') or 'Garamond'
    black = RGBColor(0, 0, 0)
    gray  = RGBColor(100, 100, 100)

    # 1. Pre-clean
    preprocess_document(doc)

    # 1b. High Priority: Hindi Unicode to Kruti Dev Conversion Pass
    # If the user selected Kruti Dev, we MUST convert the underlying character codes
    if is_krutidev(font_name):
        for para in doc.paragraphs:
            # Convert paragraph text while preserving runs if possible
            for run in para.runs:
                if run.text.strip():
                    run.text = unicode_to_krutidev(run.text)
        # Also handle tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text.strip():
                                run.text = unicode_to_krutidev(run.text)

    # 2. Page Size — thesis wider left margin for binding; letter tighter
    page_size_key = opts.get('page_size', 'A4')
    page_w, page_h = PAGE_SIZE_MAP.get(page_size_key, PAGE_SIZE_MAP['A4'])
    for section in doc.sections:
        section.page_width   = page_w
        section.page_height  = page_h
        if doc_type == 'thesis':
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)
        elif doc_type == 'letter':
            # Preserve original margins for letters with embedded letterhead/images
            # Only override if user explicitly sets page_size different from original
            if opts.get('page_size') and opts.get('page_size') != 'A4':
                section.top_margin    = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin   = Inches(1.2)
                section.right_margin  = Inches(1.0)
            # else: keep original section margins intact
        else:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)

    # 2b. Center all tables
    center_all_tables(doc)

    # 3. Title page / header — routed by doc_type
    if doc_type == 'thesis':
        insert_thesis_title_page(doc, opts, font_name)
    elif doc_type == 'letter':
        # Only inject header if user provided opts AND document has no existing Ref/Date line
        has_user_header = opts.get('org_name') or opts.get('subject')
        if has_user_header and not has_existing_letter_header(doc):
            insert_letter_header(doc, opts, font_name)
        # else: preserve original letterhead structure as-is
    else:
        insert_title_page(doc, opts, font_name)

    # 4. Body formatting — routed by doc_type
    if doc_type == 'thesis':
        format_thesis_body(doc, opts, font_name)
    elif doc_type == 'letter':
        format_letter_body(doc, opts, font_name)
    else:
        # ── BOOK / RESEARCH ──
        krutidev_mode = is_krutidev(font_name)
        black = RGBColor(0, 0, 0)
        base_size = float(opts.get('font_size', 12))
        line_spacing = float(opts.get('line_spacing', 1.15))

        i = 0
        prev_etype = None
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            text = para.text.strip()
            if not text:
                i += 1
                continue

            etype = detect_structure(para, i, doc)
            if etype == 'empty':
                i += 1
                continue
            
            # Default spacing rule: 4.0 spacing after every paragraph
            space_after = 4.0
            
            # Look ahead to see if next paragraph is body/bullet
            next_etype = None
            if i < len(doc.paragraphs) - 1:
                next_para = doc.paragraphs[i+1]
                if next_para.text.strip():
                    next_etype = detect_structure(next_para, i+1, doc)
            
            # If subheading is followed by body or bullet, reduce space_after
            if etype == 'subheading' and next_etype in ['body', 'bullet']:
                space_after = 2.0

            # Reduce space before if the previous paragraph was also a heading
            space_before = 14.0
            if etype in ['subheading'] and prev_etype in ['chapter_title', 'subheading']:
                space_before = 2.0

            if etype == 'paper_title':
                # Research paper title: centered, large, bold, no indent
                apply_para_formatting(para, etype, font_name,
                    font_size_pt=base_size + 4, bold=True, color=black,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=6, space_after_pt=10,
                    line_spacing=line_spacing)

            elif etype == 'author_name':
                # Author name: centered, medium, bold
                apply_para_formatting(para, etype, font_name,
                    font_size_pt=base_size + 1, bold=True, color=black,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=6, space_after_pt=2,
                    line_spacing=line_spacing)

            elif etype == 'author_role':
                # Affiliation/role: centered, normal
                apply_para_formatting(para, etype, font_name,
                    font_size_pt=base_size, bold=False, color=black,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=0, space_after_pt=4,
                    line_spacing=line_spacing)

            elif etype == 'title':
                apply_para_formatting(para, etype, font_name,
                    font_size_pt=base_size + 16, bold=True, color=black,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before_pt=72, space_after_pt=36,
                    line_spacing=line_spacing)

            elif etype == 'chapter_title':
                # Handle 'Chapter 8: Title' by splitting it for books too
                if ':' in text and text.lower().startswith('chapter'):
                    parts = text.split(':', 1)
                    chapter_label = parts[0].strip()
                    chapter_title = parts[1].strip()
                    
                    para.text = chapter_label.upper()
                    apply_para_formatting(para, etype, font_name,
                        font_size_pt=base_size + 8, bold=True, color=black,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before_pt=48, space_after_pt=0,
                        line_spacing=line_spacing)
                    
                    title_para = doc.add_paragraph(chapter_title)
                    para._p.addnext(title_para._p)
                    
                    apply_para_formatting(title_para, etype, font_name,
                        font_size_pt=base_size + 8, bold=True, color=black,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before_pt=0, space_after_pt=24,
                        line_spacing=line_spacing)
                    i += 2
                    prev_etype = 'chapter_title'
                    continue
                else:
                    apply_para_formatting(para, etype, font_name,
                        font_size_pt=base_size + 8, bold=True, color=black,
                        align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before_pt=48, space_after_pt=24,
                        line_spacing=line_spacing)

            elif etype == 'subheading':
                # Preserve original alignment for subheadings
                orig_align = get_original_alignment(para) or WD_ALIGN_PARAGRAPH.LEFT
                apply_para_formatting(para, etype, font_name,
                    font_size_pt=base_size + 1, bold=True, color=black,
                    align=orig_align,
                    space_before_pt=space_before, space_after_pt=6,
                    line_spacing=line_spacing)
                if not krutidev_mode and ':' in para.text:
                    apply_bold_before_colon(para, font_name, krutidev_mode)

            elif etype == 'bullet':
                is_bold_para = is_all_bold(para)
                apply_para_formatting(para, etype, font_name,
                    font_size_pt=base_size, bold=is_bold_para, color=black,
                    align=WD_ALIGN_PARAGRAPH.LEFT,
                    space_before_pt=0, space_after_pt=4,
                    line_spacing=line_spacing)
                if not krutidev_mode and ':' in para.text and not is_bold_para:
                    apply_bold_before_colon(para, font_name, krutidev_mode)

            else:  # body
                # Sentence detection logic for indent
                sentences = re.split(r'[.!?](?:\s|$)', text)
                sentences = [s for s in sentences if s.strip()]
                
                # Rule: Multiple sentences get indent, single sentence starts from margin
                indent_val = Inches(0.5) if len(sentences) > 1 else None

                apply_clean_justify(para)
                final_align = para.alignment if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY else WD_ALIGN_PARAGRAPH.LEFT

                apply_para_formatting(para, etype, font_name,
                    font_size_pt=base_size, bold=False, color=black,
                    align=final_align,
                    space_before_pt=0, space_after_pt=space_after,
                    first_indent=indent_val,
                    line_spacing=line_spacing)
                
                # Bold label before colon only for short lines
                if ':' in para.text and len(para.text.split()) <= 20:
                    apply_bold_before_colon(para, font_name, krutidev_mode)
            
            prev_etype = etype
            i += 1

    # 5. Headers & Footers
    header_text  = opts.get('header', '').strip()
    footer_text  = opts.get('footer', '').strip()
    page_numbers = opts.get('page_numbers', False)
    page_num_pos = opts.get('page_number_position', 'center')
    start_page   = opts.get('start_page_number', 1)
    try:
        start_page = int(start_page)
    except (ValueError, TypeError):
        start_page = 1
    ALIGN_MAP = {
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
    }
    num_align = ALIGN_MAP.get(page_num_pos, WD_ALIGN_PARAGRAPH.CENTER)

    for section in doc.sections:
        # Set start page number
        if page_numbers and start_page != 1:
            sectPr = section._sectPr
            pgNumType = sectPr.find(qn('w:pgNumType'))
            if pgNumType is None:
                pgNumType = OxmlElement('w:pgNumType')
                sectPr.append(pgNumType)
            pgNumType.set(qn('w:start'), str(start_page))

        if header_text:
            hdr_para = section.header.paragraphs[0]
            hdr_para.clear()
            hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = hdr_para.add_run(header_text)
            set_font_properly(r, font_name)
            r.font.size = Pt(9)
            r.font.color.rgb = gray

        ftr = section.footer
        for fp in ftr.paragraphs:
            fp.clear()

        if footer_text:
            ft_para = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
            ft_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = ft_para.add_run(footer_text)
            set_font_properly(r, font_name)
            r.font.size = Pt(9)
            r.font.color.rgb = gray

        if page_numbers:


            pn_para = ftr.add_paragraph()
            pn_para.alignment = num_align
            r1 = pn_para.add_run()
            set_font_properly(r1, font_name)
            r1.font.size = Pt(9)
            r1.font.color.rgb = gray
            add_fld_char(r1, 'begin')
            add_instr_text(r1, ' PAGE ')
            add_fld_char(r1, 'end')


    doc.save(output_file)




    #         pn_para = ftr.add_paragraph()
    #         pn_para.alignment = num_align
    #         r1 = pn_para.add_run()
    #         set_font_properly(r1, font_name)
    #         r1.font.size = Pt(9)
    #         r1.font.color.rgb = gray
    #         add_fld_char(r1, 'begin')
    #         add_instr_text(r1, ' PAGE ')
    #         add_fld_char(r1, 'end')
    #         r2 = pn_para.add_run(' / ')
    #         set_font_properly(r2, font_name)
    #         r2.font.size = Pt(9)
    #         r2.font.color.rgb = gray
    #         r3 = pn_para.add_run()
    #         set_font_properly(r3, font_name)
    #         r3.font.size = Pt(9)
    #         r3.font.color.rgb = gray
    #         add_fld_char(r3, 'begin')
    #         add_instr_text(r3, ' NUMPAGES ')
    #         add_fld_char(r3, 'end')

    # doc.save(output_file)


if __name__ == '__main__':
    in_p   = sys.argv[1]
    out_p  = sys.argv[2]
    type_d = sys.argv[3]
    opts_f = sys.argv[4]

    options = {}
    if os.path.exists(opts_f) and os.path.getsize(opts_f) > 0:
        with open(opts_f, 'r', encoding='utf-8') as f:
            options = json.load(f)

    format_document(in_p, out_p, options, doc_type=type_d)
    print(f'Success: {out_p}')




"""




"""